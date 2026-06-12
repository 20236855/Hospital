"""
医院信息管理系统 - 业务用例时序图Word文档生成工具（优化版）
使用python-docx生成专业的Word文档，包含高亮代码块
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

# 时序图数据
sequence_diagrams = [
    {
        "id": "1",
        "title": "1. 用户登录流程（通用）",
        "mermaid": '''sequenceDiagram
    participant Client as 前端
    participant Gateway as Gateway网关
    participant Auth as ruoyi-auth
    participant Redis as Redis缓存

    Client->>Gateway: POST /auth/login
    Gateway->>Auth: 转发登录请求
    Auth->>Auth: 验证用户名密码
    Auth->>Redis: 缓存用户信息
    Auth-->>Gateway: 返回LoginUser
    Gateway->>Gateway: 生成JWT Token
    Gateway-->>Client: {"access_token": "...", "token_type": "Bearer"}''',
        "description": """【用例描述】
• 功能说明：系统登录认证，支持管理员、医生、患者三种角色登录
• 触发条件：用户在登录页面输入账号密码后点击登录
• 涉及模块：ruoyi-auth（认证服务）、ruoyi-gateway（API网关）
• 核心流程：前端发起登录请求 → 网关转发 → 认证服务验证 → Redis缓存用户 → 返回JWT Token
• 业务规则：密码验证失败3次后账号锁定5分钟"""
    },
    {
        "id": "2",
        "title": "2. 患者注册流程",
        "mermaid": '''sequenceDiagram
    participant Client as 前端
    participant Gateway as Gateway网关
    participant Auth as ruoyi-auth
    participant System as ruoyi-system
    participant Patient as his-patient
    participant DB as MySQL

    Client->>Gateway: POST /auth/register
    Gateway->>Auth: 转发注册请求
    Auth->>System: 创建系统用户
    System->>DB: INSERT INTO sys_user
    Auth->>Patient: 创建患者档案
    Patient->>DB: INSERT INTO patient
    Patient-->>Auth: 返回患者信息
    Auth-->>Gateway: 注册成功
    Gateway-->>Client: {"code": 200, "msg": "注册成功"}''',
        "description": """【用例描述】
• 功能说明：患者在线注册，创建系统账号和患者档案
• 触发条件：新用户在注册页面填写个人信息后提交
• 涉及模块：ruoyi-auth、ruoyi-system、his-patient
• 核心流程：前端发起注册 → 创建系统用户 → 创建患者档案 → 返回注册成功
• 业务规则：身份证号需唯一，手机号需验证"""
    },
    {
        "id": "3",
        "title": "3. 在线挂号流程",
        "mermaid": '''sequenceDiagram
    participant Client as 患者端
    participant Gateway as Gateway网关
    participant Register as his-register
    participant Doctor as his-doctor
    participant Schedule as his-doctor
    participant DB as MySQL

    Client->>Gateway: GET /register/dept/list
    Gateway->>Register: 查询科室列表
    Register->>DB: SELECT * FROM sys_dept
    DB-->>Register: 返回科室数据
    Register-->>Gateway: 返回科室列表
    Gateway-->>Client: 科室列表

    Client->>Gateway: GET /register/doctor/list/{deptId}
    Gateway->>Register: 根据科室查医生
    Register->>Doctor: 查询医生
    Doctor->>DB: SELECT * FROM doctor WHERE deptId=?
    DB-->>Doctor: 返回医生数据
    Doctor-->>Register: 返回医生列表
    Register-->>Gateway: 返回医生列表
    Gateway-->>Client: 医生列表

    Client->>Gateway: GET /register/getFee?doctorId=1
    Gateway->>Register: 获取挂号费
    Register->>DB: SELECT fee FROM doctor_level WHERE levelId=?
    DB-->>Register: 返回费用
    Register-->>Gateway: 返回挂号费
    Gateway-->>Client: {"fee": 50.00}

    Client->>Gateway: POST /register
    Gateway->>Register: 创建挂号单
    Register->>Schedule: 扣减号源
    Schedule->>DB: UPDATE schedule SET reservedNumber=reservedNumber+1
    Register->>DB: INSERT INTO register
    DB-->>Register: 返回挂号记录
    Register-->>Gateway: 返回挂号信息
    Gateway-->>Client: {"registerNo": "RG202606110001", "status": "UNPAID"}''',
        "description": """【用例描述】
• 功能说明：患者在线选择科室、医生、排班时段完成挂号
• 触发条件：患者登录后选择挂号服务
• 涉及模块：his-register、his-doctor
• 核心流程：选择科室 → 选择医生 → 获取挂号费 → 创建挂号单 → 扣减号源
• 业务规则：同一患者同一天同一医生只能挂一个号"""
    },
    {
        "id": "4",
        "title": "4. 线下挂号流程",
        "mermaid": '''sequenceDiagram
    participant Client as 收费员端
    participant Gateway as Gateway网关
    participant Register as his-register
    participant Patient as his-patient
    participant Payment as his-payment
    participant DB as MySQL

    Client->>Gateway: GET /patient/search?idCard=xxx
    Gateway->>Patient: 查询患者
    Patient->>DB: SELECT * FROM patient WHERE idCard=?
    DB-->>Patient: 返回患者
    Patient-->>Gateway: 返回患者信息
    Gateway-->>Client: 患者信息

    Client->>Gateway: POST /register (registerType=offline)
    Gateway->>Register: 创建线下挂号单
    Register->>DB: INSERT INTO register
    DB-->>Register: 返回挂号记录
    Register->>Payment: 创建收费单
    Payment->>DB: INSERT INTO payment (payStatus=PAID)
    DB-->>Payment: 返回支付记录
    Payment-->>Register: 支付成功
    Register->>DB: UPDATE register SET payStatus=PAID
    Register-->>Gateway: 返回挂号信息
    Gateway-->>Client: {"registerNo": "RG202606110002", "payStatus": "PAID"}''',
        "description": """【用例描述】
• 功能说明：收费员为患者办理线下挂号，支持现金支付
• 触发条件：患者到医院窗口办理挂号
• 涉及模块：his-register、his-patient、his-payment
• 核心流程：查询患者 → 创建挂号单 → 创建收费单（已支付） → 更新状态
• 业务规则：线下挂号默认已支付，需打印挂号凭条"""
    },
    {
        "id": "5",
        "title": "5. 患者签到流程",
        "mermaid": '''sequenceDiagram
    participant Client as 患者端
    participant Gateway as Gateway网关
    participant CheckIn as his-register
    participant Register as his-register
    participant DB as MySQL

    Client->>Gateway: POST /register/in/checkin
    Gateway->>CheckIn: 签到请求
    CheckIn->>Register: 验证挂号单
    Register->>DB: SELECT * FROM register WHERE registerId=?
    DB-->>Register: 返回挂号记录
    alt 挂号有效且未签到
        Register-->>CheckIn: 验证通过
        CheckIn->>DB: INSERT INTO check_in
        DB-->>CheckIn: 返回签到记录
        CheckIn-->>Gateway: 签到成功
        Gateway-->>Client: {"queueNo": 3, "status": "CHECKED_IN"}
    else 挂号无效或已签到
        Register-->>CheckIn: 验证失败
        CheckIn-->>Gateway: 签到失败
        Gateway-->>Client: {"code": 500, "msg": "签到失败"}
    end''',
        "description": """【用例描述】
• 功能说明：患者到达医院后进行签到，生成排队号
• 触发条件：患者在签到机或移动端进行签到
• 涉及模块：his-register
• 核心流程：发起签到 → 验证挂号单 → 创建签到记录 → 生成排队号
• 业务规则：签到时间需在挂号时间前后30分钟内"""
    },
    {
        "id": "6",
        "title": "6. 医生接诊流程",
        "mermaid": '''sequenceDiagram
    participant Client as 医生端
    participant Gateway as Gateway网关
    participant Encounter as his-emr
    participant Register as his-register
    participant DB as MySQL

    Client->>Gateway: GET /emr/encounter/list?doctorId=1
    Gateway->>Encounter: 查询待接诊列表
    Encounter->>DB: SELECT * FROM encounter WHERE doctorId=? AND status=WAITING
    DB-->>Encounter: 返回待接诊列表
    Encounter-->>Gateway: 返回接诊列表
    Gateway-->>Client: 待接诊患者列表

    Client->>Gateway: PUT /emr/encounter/start/{encounterId}
    Gateway->>Encounter: 开始接诊
    Encounter->>DB: UPDATE encounter SET status=IN_PROGRESS
    DB-->>Encounter: 更新成功
    Encounter-->>Gateway: 接诊开始
    Gateway-->>Client: {"status": "IN_PROGRESS", "visitTime": "2026-06-11 09:00:00"}''',
        "description": """【用例描述】
• 功能说明：医生查看待接诊患者列表并开始接诊
• 触发条件：医生登录后查看接诊队列
• 涉及模块：his-emr、his-register
• 核心流程：查询待接诊列表 → 选择患者 → 开始接诊 → 更新接诊状态
• 业务规则：按排队顺序接诊，接诊状态变更记录时间"""
    },
    {
        "id": "7",
        "title": "7. 电子病历书写流程",
        "mermaid": '''sequenceDiagram
    participant Client as 医生端
    participant Gateway as Gateway网关
    participant EMR as his-emr
    participant Patient as his-patient
    participant DB as MySQL

    Client->>Gateway: GET /emr/record/{encounterId}
    Gateway->>EMR: 获取病历模板
    EMR->>Patient: 获取患者既往史
    Patient->>DB: SELECT * FROM patient WHERE patientId=?
    DB-->>Patient: 返回患者信息
    Patient-->>EMR: 返回患者既往史
    EMR-->>Gateway: 返回病历模板+既往史
    Gateway-->>Client: 病历模板

    Client->>Gateway: POST /emr/record
    Gateway->>EMR: 保存病历
    EMR->>DB: INSERT INTO medical_record
    DB-->>EMR: 返回病历ID
    EMR-->>Gateway: 保存成功
    Gateway-->>Client: {"recordId": 1, "status": "SAVED"}''',
        "description": """【用例描述】
• 功能说明：医生在接诊过程中书写电子病历
• 触发条件：医生开始接诊后进入病历书写界面
• 涉及模块：his-emr、his-patient
• 核心流程：获取病历模板 → 填充患者信息 → 书写病历内容 → 保存病历
• 业务规则：病历需包含主诉、现病史、体格检查、诊断意见、治疗方案"""
    },
    {
        "id": "8",
        "title": "8. 检查检验申请流程",
        "mermaid": '''sequenceDiagram
    participant Client as 医生端
    participant Gateway as Gateway网关
    participant Exam as his-exam
    participant EMR as his-emr
    participant DB as MySQL

    Client->>Gateway: GET /exam/technology/list
    Gateway->>Exam: 获取医技项目列表
    Exam->>DB: SELECT * FROM exam_technology
    DB-->>Exam: 返回医技项目
    Exam-->>Gateway: 返回项目列表
    Gateway-->>Client: 医技项目列表

    Client->>Gateway: POST /exam/apply
    Gateway->>Exam: 创建检查申请
    Exam->>DB: INSERT INTO exam_apply (status=UNPAID)
    DB-->>Exam: 返回申请ID
    Exam-->>EMR: 关联接诊记录
    EMR->>DB: UPDATE encounter SET examApplyId=?
    Exam-->>Gateway: 创建成功
    Gateway-->>Client: {"applyId": 1, "status": "UNPAID"}''',
        "description": """【用例描述】
• 功能说明：医生根据病情开具检查或检验申请
• 触发条件：医生在病历书写过程中需要安排检查
• 涉及模块：his-exam、his-emr
• 核心流程：选择医技项目 → 创建检查申请 → 关联接诊记录
• 业务规则：检查申请需先缴费才能执行"""
    },
    {
        "id": "9",
        "title": "9. 处方开具流程",
        "mermaid": '''sequenceDiagram
    participant Client as 医生端
    participant Gateway as Gateway网关
    participant Prescription as his-prescription
    participant EMR as his-emr
    participant DB as MySQL

    Client->>Gateway: GET /prescription/drug/list
    Gateway->>Prescription: 获取药品列表
    Prescription->>DB: SELECT * FROM drug
    DB-->>Prescription: 返回药品
    Prescription-->>Gateway: 返回药品列表
    Gateway-->>Client: 药品列表

    Client->>Gateway: POST /prescription/item
    Gateway->>Prescription: 创建处方明细
    Prescription->>DB: INSERT INTO prescription_item
    DB-->>Prescription: 返回明细ID
    Prescription-->>EMR: 关联病历
    EMR->>DB: UPDATE medical_record SET prescriptionId=?
    Prescription-->>Gateway: 创建成功
    Gateway-->>Client: {"itemId": 1, "status": "PENDING"}''',
        "description": """【用例描述】
• 功能说明：医生根据诊断结果开具药品处方
• 触发条件：医生完成诊断后需要开具处方
• 涉及模块：his-prescription、his-emr
• 核心流程：选择药品 → 创建处方明细 → 关联病历
• 业务规则：处方需包含药品名称、剂量、用法用量"""
    },
    {
        "id": "10",
        "title": "10. 收费支付流程",
        "mermaid": '''sequenceDiagram
    participant Client as 患者端
    participant Gateway as Gateway网关
    participant Payment as his-payment
    participant Register as his-register
    participant Exam as his-exam
    participant DB as MySQL

    Client->>Gateway: GET /payment/amount?registerId=1
    Gateway->>Payment: 计算总费用
    Payment->>Register: 获取挂号费用
    Register->>DB: SELECT registerFee FROM register WHERE registerId=?
    DB-->>Register: 返回费用
    Register-->>Payment: 返回挂号费
    Payment->>Exam: 获取检查费用
    Exam->>DB: SELECT SUM(price) FROM exam_apply WHERE registerId=?
    DB-->>Exam: 返回费用
    Exam-->>Payment: 返回检查费
    Payment-->>Gateway: 返回总费用
    Gateway-->>Client: {"totalAmount": 250.00}

    Client->>Gateway: POST /payment/pay
    Gateway->>Payment: 支付请求
    Payment->>DB: INSERT INTO payment
    DB-->>Payment: 返回支付记录
    Payment->>Register: 标记挂号已支付
    Register->>DB: UPDATE register SET payStatus=PAID
    Payment->>Exam: 标记检查已支付
    Exam->>DB: UPDATE exam_apply SET payStatus=PAID
    Payment-->>Gateway: 支付成功
    Gateway-->>Client: {"payNo": "PY202606110001", "status": "PAID"}''',
        "description": """【用例描述】
• 功能说明：患者完成挂号费、检查费、药费等费用的支付
• 触发条件：患者收到费用通知后进行支付
• 涉及模块：his-payment、his-register、his-exam
• 核心流程：计算总费用 → 选择支付方式 → 创建支付记录 → 更新相关单据状态
• 业务规则：支持现金、微信、支付宝、医保等支付方式"""
    },
    {
        "id": "11",
        "title": "11. 检查报告查看流程",
        "mermaid": '''sequenceDiagram
    participant Client as 患者端/医生端
    participant Gateway as Gateway网关
    participant Exam as his-exam
    participant DB as MySQL

    Client->>Gateway: GET /exam/apply/list?patientId=1
    Gateway->>Exam: 查询检查申请
    Exam->>DB: SELECT * FROM exam_apply WHERE patientId=?
    DB-->>Exam: 返回申请列表
    Exam-->>Gateway: 返回申请列表
    Gateway-->>Client: 检查申请列表

    Client->>Gateway: GET /exam/apply/{applyId}
    Gateway->>Exam: 获取检查详情
    Exam->>DB: SELECT * FROM exam_apply WHERE applyId=?
    DB-->>Exam: 返回申请详情
    Exam-->>Gateway: 返回详情(含报告)
    Gateway-->>Client: {"examResult": "...", "imageUrl": "xxx"}''',
        "description": """【用例描述】
• 功能说明：患者或医生查看检查报告和影像结果
• 触发条件：检查完成后患者或医生查看报告
• 涉及模块：his-exam
• 核心流程：查询检查列表 → 选择检查项 → 查看报告详情
• 业务规则：患者只能查看自己的报告，医生可查看自己患者的报告"""
    },
    {
        "id": "12",
        "title": "12. 退号流程",
        "mermaid": '''sequenceDiagram
    participant Client as 患者端/收费员端
    participant Gateway as Gateway网关
    participant Register as his-register
    participant Schedule as his-doctor
    participant Payment as his-payment
    participant DB as MySQL

    Client->>Gateway: PUT /register/cancel/{registerId}
    Gateway->>Register: 退号请求
    Register->>DB: SELECT * FROM register WHERE registerId=?
    DB-->>Register: 返回挂号记录
    alt 可退号(未就诊、未检查)
        Register->>Schedule: 恢复号源
        Schedule->>DB: UPDATE schedule SET reservedNumber=reservedNumber-1
        Register->>Payment: 发起退款
        Payment->>DB: UPDATE payment SET payStatus=REFUND
        Register->>DB: UPDATE register SET registerStatus=CANCELLED
        Register-->>Gateway: 退号成功
        Gateway-->>Client: {"status": "CANCELLED", "refundAmount": 50.00}
    else 不可退号
        Register-->>Gateway: 退号失败
        Gateway-->>Client: {"code": 500, "msg": "已就诊，无法退号"}
    end''',
        "description": """【用例描述】
• 功能说明：患者申请退号，系统处理退款并恢复号源
• 触发条件：患者在就诊前申请退号
• 涉及模块：his-register、his-doctor、his-payment
• 核心流程：发起退号请求 → 验证可退条件 → 恢复号源 → 发起退款 → 更新状态
• 业务规则：已就诊或已检查的挂号单不可退号"""
    }
]


def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_code_block(doc, mermaid_code):
    """添加带边框的代码块"""
    # 添加代码块标签
    label_para = doc.add_paragraph()
    label_run = label_para.add_run('【Mermaid时序图代码】')
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 创建带边框的段落
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.3)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    
    # 设置段落背景色
    pPr = code_para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')  # 浅灰色背景
    pPr.append(shd)
    
    # 添加代码文本
    code_run = code_para.add_run(mermaid_code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = RGBColor(33, 33, 33)
    
    doc.add_paragraph()


def create_word_document():
    """创建Word文档"""
    print("正在创建Word文档...")
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    
    # ==================== 文档标题 ====================
    title = doc.add_heading('', 0)
    title_run = title.add_run('医院信息管理系统')
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('', level=1)
    subtitle_run = subtitle.add_run('业务用例时序图文档')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(51, 51, 51)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分隔线
    doc.add_paragraph()
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run('━' * 60)
    line_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # 添加文档信息
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_run1 = info_para.add_run('项目名称：医院信息管理系统（HIS）')
    info_run1.font.size = Pt(11)
    info_para.add_run('\n')
    
    info_run2 = info_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    info_run2.font.size = Pt(11)
    info_para.add_run('\n')
    
    info_run3 = info_para.add_run('文档版本：V1.0')
    info_run3.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ==================== 目录 ====================
    toc_heading = doc.add_heading('', level=1)
    toc_run = toc_heading.add_run('目  录')
    toc_run.font.color.rgb = RGBColor(0, 51, 102)
    
    for i, diagram in enumerate(sequence_diagrams, 1):
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.left_indent = Inches(0.5)
        toc_para.paragraph_format.space_after = Pt(8)
        
        num_run = toc_para.add_run(f'{i}. ')
        num_run.bold = True
        num_run.font.color.rgb = RGBColor(0, 102, 204)
        
        title_run = toc_para.add_run(diagram['title'])
        title_run.font.color.rgb = RGBColor(51, 51, 51)
    
    doc.add_page_break()
    
    # ==================== 用例内容 ====================
    for i, diagram in enumerate(sequence_diagrams, 1):
        print(f"正在处理 {diagram['title']}...")
        
        # 章节标题
        heading = doc.add_heading('', level=1)
        heading_run = heading.add_run(diagram['title'])
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(12)
        
        # 添加代码块
        add_code_block(doc, diagram['mermaid'])
        
        # 添加用例描述
        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.left_indent = Inches(0.3)
        desc_para.paragraph_format.space_after = Pt(12)
        
        # 处理描述文本
        lines = diagram['description'].strip().split('\n')
        for j, line in enumerate(lines):
            if j == 0:
                # 标题行加粗
                title_run = desc_para.add_run(line)
                title_run.bold = True
                title_run.font.size = Pt(11)
                title_run.font.color.rgb = RGBColor(0, 102, 153)
            else:
                # 内容行
                content_run = desc_para.add_run(line)
                content_run.font.size = Pt(10)
                content_run.font.color.rgb = RGBColor(51, 51, 51)
            desc_para.add_run('\n')
        
        # 添加分隔线
        doc.add_paragraph()
        sep_para = doc.add_paragraph()
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = sep_para.add_run('─' * 70)
        sep_run.font.color.rgb = RGBColor(200, 200, 200)
        
        # 每3个用例后添加分页符
        if i % 3 == 0 and i < len(sequence_diagrams):
            doc.add_page_break()
    
    # ==================== 附录 ====================
    doc.add_page_break()
    
    appendix_heading = doc.add_heading('', level=1)
    appendix_run = appendix_heading.add_run('附录：系统模块与数据字典')
    appendix_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 模块说明
    module_heading = doc.add_heading('', level=2)
    module_run = module_heading.add_run('A. 系统模块说明')
    module_run.font.color.rgb = RGBColor(51, 102, 153)
    
    modules = [
        ('ruoyi-auth', '认证授权服务', '处理用户登录、注册、Token管理'),
        ('ruoyi-gateway', 'API网关', '请求路由、过滤器、负载均衡'),
        ('ruoyi-system', '系统管理', '用户管理、角色权限、菜单管理'),
        ('his-patient', '患者管理', '患者档案、信息维护'),
        ('his-register', '挂号管理', '在线/线下挂号、签到管理'),
        ('his-doctor', '医生管理', '医生信息、排班管理'),
        ('his-emr', '电子病历', '接诊管理、病历书写'),
        ('his-exam', '检查检验', '检查申请、报告管理'),
        ('his-payment', '收费管理', '费用计算、支付、退费'),
        ('his-prescription', '处方管理', '处方开具、药品管理'),
    ]
    
    # 创建表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # 设置列宽
    for i, width in enumerate([Inches(1.8), Inches(2.0), Inches(3.2)]):
        table.columns[i].width = width
    
    # 表头
    header_cells = table.rows[0].cells
    headers = ['模块名称', '功能描述', '说明']
    for j, header in enumerate(headers):
        header_cells[j].text = header
        header_cells[j].paragraphs[0].runs[0].font.bold = True
        header_cells[j].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(header_cells[j], 'D9E2F3')
    
    # 数据行
    for module_name, desc, note in modules:
        row_cells = table.add_row().cells
        row_cells[0].text = module_name
        row_cells[1].text = desc
        row_cells[2].text = note
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # 数据字典
    table_heading = doc.add_heading('', level=2)
    table_run = table_heading.add_run('B. 核心数据表说明')
    table_run.font.color.rgb = RGBColor(51, 102, 153)
    
    tables = [
        ('patient', '患者信息表', '存储患者基本信息'),
        ('register', '挂号记录表', '存储挂号信息'),
        ('check_in', '签到记录表', '存储签到信息'),
        ('doctor', '医生信息表', '存储医生信息'),
        ('schedule', '排班表', '存储医生排班'),
        ('encounter', '接诊记录表', '存储接诊信息'),
        ('medical_record', '电子病历表', '存储病历信息'),
        ('exam_apply', '检查申请表', '存储检查申请'),
        ('prescription', '处方表', '存储处方信息'),
        ('payment', '支付记录表', '存储支付信息'),
    ]
    
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    
    for i, width in enumerate([Inches(1.8), Inches(2.0), Inches(3.2)]):
        table2.columns[i].width = width
    
    header_cells2 = table2.rows[0].cells
    headers2 = ['表名', '中文名称', '说明']
    for j, header in enumerate(headers2):
        header_cells2[j].text = header
        header_cells2[j].paragraphs[0].runs[0].font.bold = True
        header_cells2[j].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(header_cells2[j], 'E2EFDA')
    
    for table_name, chinese_name, note in tables:
        row_cells = table2.add_row().cells
        row_cells[0].text = table_name
        row_cells[1].text = chinese_name
        row_cells[2].text = note
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # 添加说明
    note_para = doc.add_paragraph()
    note_run = note_para.add_run('※ 文档说明：')
    note_run.bold = True
    note_run.font.color.rgb = RGBColor(153, 0, 0)
    
    note_para.add_run('本文档包含12个核心业务用例的时序图Mermaid代码。您可以使用Mermaid Live Editor（')
    note_para.add_run('https://mermaid.live/').font.color.rgb = RGBColor(0, 102, 204)
    note_para.add_run('）打开代码并导出为PNG图片，或在支持Mermaid的编辑器中实时预览。')
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(__file__), '医院信息管理系统_业务用例时序图.docx')
    doc.save(output_path)
    
    print(f"\n{'='*60}")
    print(f"✅ Word文档已成功生成！")
    print(f"{'='*60}")
    print(f"📄 文档路径：{output_path}")
    print(f"� 包含内容：")
    print(f"   • 12个业务用例时序图（Mermaid代码）")
    print(f"   • 每个用例的详细功能说明")
    print(f"   • 系统模块说明表")
    print(f"   • 核心数据表说明")
    print(f"{'='*60}")
    
    return output_path


if __name__ == '__main__':
    create_word_document()
